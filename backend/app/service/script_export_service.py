"""ScriptLens · 剧本完整导出服务（F3）。

职责：
  - 拼装一份"应用了全部历史改写后的最终剧本全文"
  - 以 docx / pdf / txt 三种格式渲染并返回 bytes

应用最新改写的策略（与 script_operations 表对齐）：
  对每个 scene，找它最近一条 success=True 的 rewrite op；
  - 有：用 snapshot_after[scene_id] 作为该场最终文本
  - 无：用 scenes.text 原始文本

故意不做：
  - 没有"接受/拒绝改写"的概念。当前后端只有"记录改写"（不写回 scenes.text），
    所以"最新成功改写"就是用户在 UI 看到的最新版本（PRD §8）。如果将来上 scene_versions
    表，这里只需要换数据源即可。

字体：
  reportlab 的 STSong-Light 是内置 CID 字体，对简中支持完整、不依赖系统字体文件。
"""

from __future__ import annotations

import io
import logging
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from sqlalchemy import text
from sqlalchemy.engine import Engine

from utils.database import engine as default_engine

logger = logging.getLogger(__name__)


# 注册 STSong-Light CID 字体（reportlab 内置，无需额外字体文件）
# 多次注册是安全的，但避免在请求路径上反复跑：在模块加载时 try 一次。
try:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    _PDF_FONT = "STSong-Light"
except Exception as exc:  # noqa: BLE001 — 容错并退化到 Helvetica
    logger.warning("STSong-Light 注册失败，PDF 中文可能乱码：%s", exc)
    _PDF_FONT = "Helvetica"


SUPPORTED_FORMATS = ("docx", "pdf", "txt")


class ExportError(Exception):
    """导出失败（剧本不存在 / 无权 / 无场次 / 渲染失败）。"""


class ScriptNotFoundError(ExportError):
    pass


class ScriptPermissionError(ExportError):
    pass


# ============================================================
# 数据装配
# ============================================================


def _load_script_meta(
    *, script_id: str, user_id: int, conn,
) -> Tuple[str, str]:
    """返回 (title, source_format)；越权抛 ScriptPermissionError。"""
    row = conn.execute(
        text(
            """
            SELECT user_id, title, COALESCE(source_format, 'docx') AS source_format
            FROM scriptlens.scripts
            WHERE id = :sid
            """
        ),
        {"sid": script_id},
    ).mappings().first()
    if row is None:
        raise ScriptNotFoundError("剧本不存在")
    if int(row["user_id"]) != int(user_id):
        raise ScriptPermissionError("无权导出该剧本")
    return str(row["title"] or "未命名剧本"), str(row["source_format"])


def _load_scenes_with_latest_rewrite(
    *, script_id: str, conn,
) -> List[Dict[str, Any]]:
    """按集场顺序返回所有场次，text 字段已替换为最新成功改写版本（如有）。

    用一次 SQL 拉全部 scene + LATERAL 子查询取每场的 latest rewrite snapshot_after。
    没用 N 次往返查询，避免 50+ 场剧本退化成 50+ 次 query。
    """
    rows = conn.execute(
        text(
            """
            SELECT
                s.id::text AS scene_id,
                s.episode_no,
                s.scene_no,
                s.scene_label,
                s.text AS original_text,
                latest_op.rewritten_text
            FROM scriptlens.scenes s
            LEFT JOIN LATERAL (
                SELECT op.snapshot_after ->> s.id::text AS rewritten_text
                FROM scriptlens.script_operations op
                WHERE op.script_id = s.script_id
                  AND op.intent_type = 'rewrite'
                  AND op.success = TRUE
                  AND (op.snapshot_after ? s.id::text)
                ORDER BY op.created_at DESC
                LIMIT 1
            ) AS latest_op ON TRUE
            WHERE s.script_id = :sid
            ORDER BY s.episode_no NULLS LAST, s.scene_no, s.start_line
            """
        ),
        {"sid": script_id},
    ).mappings().all()

    out: List[Dict[str, Any]] = []
    for r in rows:
        rewritten = r.get("rewritten_text")
        final_text = (
            str(rewritten).strip()
            if rewritten and str(rewritten).strip()
            else str(r["original_text"] or "")
        )
        out.append(
            {
                "scene_id": r["scene_id"],
                "episode_no": r["episode_no"],
                "scene_no": r["scene_no"],
                "scene_label": r["scene_label"] or "",
                "text": final_text,
                "is_rewritten": bool(rewritten and str(rewritten).strip()),
            }
        )
    return out


def assemble_full_script(
    *, script_id: str, user_id: int, engine: Engine = default_engine,
) -> Dict[str, Any]:
    """拼装最终全文 + 元信息，供 3 种渲染共用。"""
    with engine.connect() as conn:
        title, source_format = _load_script_meta(
            script_id=script_id, user_id=user_id, conn=conn
        )
        scenes = _load_scenes_with_latest_rewrite(script_id=script_id, conn=conn)
    if not scenes:
        raise ExportError("剧本暂无场次，无法导出")
    return {
        "title": title,
        "source_format": source_format,
        "scenes": scenes,
        "rewritten_count": sum(1 for s in scenes if s["is_rewritten"]),
    }


# ============================================================
# 渲染：场次标题（3 种格式共用文案）
# ============================================================


def _scene_heading(scene: Dict[str, Any]) -> str:
    parts: List[str] = []
    ep = scene.get("episode_no")
    if ep is not None:
        parts.append(f"第 {ep} 集")
    sn = scene.get("scene_no")
    if sn is not None:
        parts.append(f"第 {sn} 场")
    label = scene.get("scene_label")
    if label:
        parts.append(f"《{label}》")
    if scene.get("is_rewritten"):
        parts.append("（已改写）")
    return " · ".join(parts) if parts else "未命名场次"


# ============================================================
# 三种格式渲染
# ============================================================


def render_txt(payload: Dict[str, Any]) -> bytes:
    lines: List[str] = []
    lines.append(payload["title"])
    lines.append("=" * len(payload["title"]) * 2)
    if payload["rewritten_count"]:
        lines.append(
            f"（本剧本含 {payload['rewritten_count']} 场 AI 改写，已应用最新版本）"
        )
    lines.append("")
    for scene in payload["scenes"]:
        lines.append(_scene_heading(scene))
        lines.append("-" * 40)
        lines.append(scene["text"].rstrip())
        lines.append("")
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def render_docx(payload: Dict[str, Any]) -> bytes:
    doc = Document()

    title_p = doc.add_heading(payload["title"], level=0)
    for run in title_p.runs:
        run.font.size = Pt(20)

    if payload["rewritten_count"]:
        note = doc.add_paragraph(
            f"本剧本含 {payload['rewritten_count']} 场 AI 改写，导出版本已应用最新内容。"
        )
        for run in note.runs:
            run.italic = True
            run.font.size = Pt(10)

    for scene in payload["scenes"]:
        heading = doc.add_heading(_scene_heading(scene), level=2)
        for run in heading.runs:
            run.font.size = Pt(14)
        for para_text in scene["text"].split("\n"):
            stripped = para_text.rstrip()
            if not stripped:
                continue
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_pdf(payload: Dict[str, Any]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        title=payload["title"],
    )
    base = getSampleStyleSheet()["Normal"]
    title_style = ParagraphStyle(
        "ScriptTitle",
        parent=base,
        fontName=_PDF_FONT,
        fontSize=22,
        leading=28,
        spaceAfter=12,
    )
    note_style = ParagraphStyle(
        "ScriptNote",
        parent=base,
        fontName=_PDF_FONT,
        fontSize=10,
        leading=14,
        textColor="#888888",
        spaceAfter=10,
    )
    heading_style = ParagraphStyle(
        "SceneHeading",
        parent=base,
        fontName=_PDF_FONT,
        fontSize=14,
        leading=20,
        spaceBefore=12,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "SceneBody",
        parent=base,
        fontName=_PDF_FONT,
        fontSize=11,
        leading=18,
        spaceAfter=4,
    )

    story: List[Any] = [Paragraph(_pdf_escape(payload["title"]), title_style)]
    if payload["rewritten_count"]:
        story.append(
            Paragraph(
                _pdf_escape(
                    f"本剧本含 {payload['rewritten_count']} 场 AI 改写，导出版本已应用最新内容。"
                ),
                note_style,
            )
        )

    for scene in payload["scenes"]:
        story.append(Paragraph(_pdf_escape(_scene_heading(scene)), heading_style))
        for para_text in scene["text"].split("\n"):
            stripped = para_text.rstrip()
            if not stripped:
                story.append(Spacer(1, 4))
                continue
            story.append(Paragraph(_pdf_escape(stripped), body_style))
        story.append(Spacer(1, 8))

    doc.build(story)
    return buf.getvalue()


def _pdf_escape(s: str) -> str:
    """reportlab Paragraph 用 HTML-like 标记，需要转义 & < >。"""
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


# ============================================================
# 公共入口
# ============================================================


def render_export(
    *,
    script_id: str,
    user_id: int,
    fmt: str,
    engine: Engine = default_engine,
) -> Tuple[bytes, str, str]:
    """导出主入口：返回 (文件 bytes, content-type, suggested filename)。"""
    if fmt not in SUPPORTED_FORMATS:
        raise ExportError(
            f"不支持的导出格式：{fmt}；支持 {', '.join(SUPPORTED_FORMATS)}"
        )
    payload = assemble_full_script(
        script_id=script_id, user_id=user_id, engine=engine
    )
    safe_title = _safe_filename(payload["title"])
    if fmt == "txt":
        return render_txt(payload), "text/plain; charset=utf-8", f"{safe_title}.txt"
    if fmt == "docx":
        return (
            render_docx(payload),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{safe_title}.docx",
        )
    return render_pdf(payload), "application/pdf", f"{safe_title}.pdf"


def _safe_filename(name: str) -> str:
    """去掉文件名里 OS / HTTP header 不允许的字符；保留中文。"""
    cleaned = "".join(
        ch for ch in (name or "").strip() if ch not in '<>:"/\\|?*\r\n\t'
    )
    return cleaned or "script"
